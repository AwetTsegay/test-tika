from docx import Document
import os
import comtypes.client as cc

document = Document()
document.add_heading('Document with Embedded JSON Files', level=1)
document.add_paragraph('This document contains five embedded JSON files.')

document.save('data_sample/example.docx')


def embed_json_as_ole(docx_path):
    word = cc.CreateObject('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(os.path.abspath(docx_path))

    for i in range(1, 6):
        json_file = os.path.abspath(f"json_files/data_{i}.json")
        doc.Paragraphs(doc.Paragraphs.Count).Range.InsertAfter(f'\nEmbedded JSON File {i}: ')

        # Embed the JSON file as an OLE object (shown as an icon)
        doc.InlineShapes.AddOLEObject(
            FileName=json_file,
            LinkToFile=False,
            DisplayAsIcon=True,
            IconFileName='C:\\WINDOWS\\system32\\packager.dll', # Generic icon
            IconLabel=f'data_{i}.json'
        )

    doc.SaveAs(os.path.abspath('data_sample/example.docx'))
    doc.Close()
    word.Quit()

# Call the function with the path to the base document
embed_json_as_ole('data_sample/example.docx')


